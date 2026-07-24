"""Generate aligned report figures, appendices, diagrams, and the final DOCX."""

from __future__ import annotations

import sys
from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from src.data import (  # noqa: E402
    load_component_history,
    load_official_prices,
    load_prediction_dataset,
    load_sources,
)
from src.modeling import COMPONENT_FEATURES, evaluate_latest_cycle  # noqa: E402
from src.pricing import reconstruction_audit  # noqa: E402

DOCS = ROOT / "docs"
APPENDICES = ROOT / "appendices"
CHARTS = ROOT / "outputs" / "charts"
DIAGRAMS = ROOT / "outputs" / "diagrams"
REPORT_PATH = DOCS / "Ryan_Final_Project_Report.docx"

TITLE = (
    "DESIGN AND IMPLEMENTATION OF A COMPONENT-BASED FUEL PRICE "
    "PREDICTION SYSTEM USING MULTIPLE LINEAR REGRESSION IN NAIROBI, KENYA"
)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(90, 90, 90)
HEADER_FILL = "E8EEF5"

COMPONENT_LABELS = {
    "Landed_Cost": "Landed cost",
    "Distribution_Storage": "Distribution and storage",
    "Margins": "Margins",
    "Stabilization_Adjustment": "Stabilization adjustment",
    "Taxes_Levies": "Taxes and levies",
    "Fuel_Diesel": "Diesel fuel-type effect",
    "Fuel_Kerosene": "Kerosene fuel-type effect",
}


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def add_table(doc: Document, headers: list[str], rows: list[list[object]], widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = str(header)
        set_cell_shading(cell, HEADER_FILL)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)

    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for run in cells[index].paragraphs[0].runs:
                run.font.size = Pt(9)
            set_cell_margins(cells[index])

    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
    return table


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(
    doc: Document,
    text: str,
    *,
    italic: bool = False,
    align=None,
) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.italic = italic
    if align is not None:
        paragraph.alignment = align


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_figure(doc: Document, path: Path, caption: str, width=6.1) -> None:
    if not path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture = paragraph.add_run().add_picture(str(path), width=Inches(width))
    picture._inline.docPr.set("descr", caption)
    caption_paragraph = doc.add_paragraph(caption)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.style = "Caption"


def apply_report_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("MafutaPlan | BSc Information Technology Project")
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = GRAY


def add_cover(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()
    institution = doc.add_paragraph()
    institution.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = institution.add_run(
        "JOMO KENYATTA UNIVERSITY OF AGRICULTURE AND TECHNOLOGY"
    )
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = DARK_BLUE

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(36)
    title.paragraph_format.space_after = Pt(20)
    run = title.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(20, 55, 85)

    product = doc.add_paragraph()
    product.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = product.add_run("MAFUTAPLAN")
    run.bold = True
    run.font.size = Pt(17)
    run.font.color.rgb = BLUE

    for text in (
        "RYAN ALFRED NYAMBATI",
        "SCT222-0195/2021",
        "",
        "A project report submitted in partial fulfilment of the requirements "
        "for the award of a Bachelor's degree at Jomo Kenyatta University of "
        "Agriculture and Technology.",
        "",
        "JULY 2026",
    ):
        paragraph = doc.add_paragraph(text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(6)
    doc.add_page_break()


def build_figures(components: pd.DataFrame, evaluation) -> None:
    CHARTS.mkdir(parents=True, exist_ok=True)
    DIAGRAMS.mkdir(parents=True, exist_ok=True)

    means = components.groupby("Effective_From")[COMPONENT_FEATURES].mean()
    fig, ax = plt.subplots(figsize=(9, 4.8))
    for column in COMPONENT_FEATURES:
        ax.plot(means.index, means[column], marker="o", label=COMPONENT_LABELS[column])
    ax.set_title("Average verified component values by cycle")
    ax.set_ylabel("KSh per litre")
    ax.legend(ncol=2, fontsize=8)
    ax.grid(alpha=0.2)
    fig.autofmt_xdate()
    fig.tight_layout()
    fig.savefig(CHARTS / "figure_3_1_component_history.png", dpi=180)
    plt.close(fig)

    result = evaluation.results.set_index("Fuel")
    fig, ax = plt.subplots(figsize=(8, 4.8))
    result[["Target_Retail_Price", "Predicted_Retail_Price"]].plot.bar(ax=ax)
    ax.set_title("Chronological test: actual versus predicted")
    ax.set_ylabel("KSh per litre")
    ax.set_xlabel("")
    ax.legend(["Actual", "Predicted"])
    ax.tick_params(axis="x", rotation=0)
    fig.tight_layout()
    fig.savefig(CHARTS / "figure_4_1_actual_vs_predicted.png", dpi=180)
    plt.close(fig)

    coefficients = evaluation.coefficients.iloc[1:].copy()
    coefficients["Term"] = coefficients["Term"].map(COMPONENT_LABELS).fillna(
        coefficients["Term"]
    )
    fig, ax = plt.subplots(figsize=(8, 4.8))
    colors = [
        "#2E74B5" if value >= 0 else "#B24C4C"
        for value in coefficients["Coefficient"]
    ]
    ax.barh(coefficients["Term"], coefficients["Coefficient"], color=colors)
    ax.axvline(0, color="#333333", linewidth=0.8)
    ax.set_title("Coefficients learned from the chronological training sample")
    ax.set_xlabel("Fitted coefficient")
    fig.tight_layout()
    fig.savefig(CHARTS / "figure_4_2_coefficients.png", dpi=180)
    plt.close(fig)

    make_flow_diagram(
        DIAGRAMS / "system_architecture_diagram.png",
        [
            "Official EPRA\nsources",
            "Verified component\ndataset",
            "Cleaning and\nvalidation",
            "Multiple linear\nregression",
            "July 2026\navailability gate",
            "Actual-versus-\npredicted evaluation",
            "Streamlit\ninterface",
        ],
        "MafutaPlan system architecture",
    )
    make_flow_diagram(
        DIAGRAMS / "conceptual_framework.png",
        [
            "Pre-target cost\ncomponents",
            "Fuel type",
            "Multiple linear\nregression",
            "Following-cycle\nretail price",
            "Budget and journey\nplanning",
        ],
        "Conceptual framework",
    )
    make_use_case_diagram(DIAGRAMS / "use_case_diagram.png")


def make_flow_diagram(path: Path, labels: list[str], title: str) -> None:
    fig, ax = plt.subplots(figsize=(12, 3.2))
    ax.axis("off")
    xs = [0.06 + index * (0.88 / (len(labels) - 1)) for index in range(len(labels))]
    for index, (x, label) in enumerate(zip(xs, labels)):
        ax.text(
            x,
            0.5,
            label,
            ha="center",
            va="center",
            fontsize=9,
            bbox={"boxstyle": "round,pad=0.5", "fc": "#E8EEF5", "ec": "#2E74B5"},
        )
        if index < len(labels) - 1:
            ax.annotate(
                "",
                xy=(xs[index + 1] - 0.055, 0.5),
                xytext=(x + 0.055, 0.5),
                arrowprops={"arrowstyle": "->", "color": "#1F4D78", "lw": 1.5},
            )
    ax.set_title(title, fontsize=14, weight="bold")
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)


def make_use_case_diagram(path: Path) -> None:
    use_cases = [
        "View July prediction status",
        "View factors affecting price",
        "Reconstruct historical price",
        "Calculate fuel purchase cost",
        "Calculate journey cost",
        "View data and methodology",
    ]
    fig, ax = plt.subplots(figsize=(8.5, 6))
    ax.axis("off")
    ax.text(
        0.12,
        0.5,
        "Brian\nNairobi ride-hailing driver",
        ha="center",
        va="center",
        bbox={"boxstyle": "round,pad=0.6", "fc": "#FFF4D6", "ec": "#8A6D1D"},
    )
    for index, label in enumerate(use_cases):
        y = 0.88 - index * 0.145
        ax.text(
            0.68,
            y,
            label,
            ha="center",
            va="center",
            bbox={"boxstyle": "round,pad=0.5", "fc": "#E8EEF5", "ec": "#2E74B5"},
        )
        ax.annotate(
            "",
            xy=(0.52, y),
            xytext=(0.23, 0.5),
            arrowprops={"arrowstyle": "-", "color": "#607D8B"},
        )
    ax.set_title("MafutaPlan use-case diagram", fontsize=14, weight="bold")
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)


def write_appendices(
    prediction_data: pd.DataFrame,
    evaluation,
) -> None:
    APPENDICES.mkdir(parents=True, exist_ok=True)
    result = evaluation.results.copy()
    result["Model"] = "Multiple Linear Regression"
    result["MAE"] = evaluation.mae
    result["RMSE"] = evaluation.rmse
    result[
        [
            "Fuel",
            "Model",
            "Input_Cycle",
            "Target_Cycle",
            "Target_Retail_Price",
            "Predicted_Retail_Price",
            "Absolute_Error",
            "Percentage_Error",
            "MAE",
            "RMSE",
        ]
    ].to_csv(APPENDICES / "Model_Metrics.csv", index=False, float_format="%.6f")

    statistics = prediction_data[
        [*COMPONENT_FEATURES, "Target_Retail_Price"]
    ].describe().T
    statistics.reset_index(names="Variable").to_csv(
        APPENDICES / "Descriptive_Statistics.csv",
        index=False,
        float_format="%.3f",
    )

    dictionary_rows = [
        ["Input_Cycle", "Date", "month", "Cycle supplying model components", "Must precede target"],
        ["Target_Cycle", "Date", "month", "Following retail-price cycle", "Must follow input"],
        ["Fuel", "Text", "", "Super Petrol, Diesel or Kerosene", "Allowed values only"],
        ["Landed_Cost", "Decimal", "KSh/L", "Aggregate imported-product landed value", "Required"],
        ["Distribution_Storage", "Decimal", "KSh/L", "Aggregate distribution and storage costs", "Required"],
        ["Margins", "Decimal", "KSh/L", "Aggregate approved margins", "Required"],
        ["Stabilization_Adjustment", "Decimal", "KSh/L", "Signed stabilization or reconciliation", "Required"],
        ["Taxes_Levies", "Decimal", "KSh/L", "Aggregate taxes and statutory levies", "Required"],
        ["Target_Retail_Price", "Decimal", "KSh/L", "Following-cycle Nairobi maximum retail price", "Positive"],
        ["Source_ID", "Text", "", "Foreign key to official source register", "Registered HTTPS source"],
        ["Verification_Status", "Text", "", "Evidence review state", "Non-blank"],
    ]
    pd.DataFrame(
        dictionary_rows, columns=["Field", "Type", "Unit", "Meaning", "Validation"]
    ).to_csv(APPENDICES / "Data_Extraction_Sheet.csv", index=False)

    prediction_data.head(9).to_csv(
        APPENDICES / "Sample_Dataset.csv", index=False, float_format="%.2f"
    )


def build_report(
    components: pd.DataFrame,
    prediction_data: pd.DataFrame,
    sources: pd.DataFrame,
    official: pd.DataFrame,
    evaluation,
) -> Document:
    doc = Document()
    apply_report_styles(doc)
    add_cover(doc)

    add_heading(doc, "DECLARATION", 1)
    add_para(
        doc,
        "I declare that this project report is my original work and has not been "
        "submitted to another university for an academic award. All sources used "
        "have been acknowledged.",
    )
    add_para(doc, "Student: Ryan Alfred Nyambati   Signature: __________   Date: __________")
    add_para(doc, "Supervisor: __________________   Signature: __________   Date: __________")

    add_heading(doc, "ACKNOWLEDGEMENT", 1)
    add_para(
        doc,
        "I thank God, my family, supervisor, lecturers, and classmates for their "
        "support. I acknowledge the Energy and Petroleum Regulatory Authority "
        "(EPRA) for publishing the official records used in this academic project.",
    )

    add_heading(doc, "ABSTRACT", 1)
    add_para(
        doc,
        "Fuel-price changes directly affect Nairobi transport operators and small "
        "businesses. This project designed and implemented MafutaPlan, a "
        "component-based system that separates machine-learning prediction from "
        "deterministic price reconstruction, scenario analysis, and fuel-cost "
        "calculation. A pooled multiple linear regression model uses landed cost, "
        "distribution and storage, margins, stabilization adjustment, taxes and "
        "levies, and encoded fuel type to estimate the following cycle's retail "
        "price. The reviewed data contains 33 fuel-cycle records across 11 official "
        "EPRA component cycles. Chronological evaluation trains on 30 records with "
        f"targets from {evaluation.training_start:%B %Y} to "
        f"{evaluation.training_end:%B %Y} and is evaluated on three April 2026 records. "
        f"The holdout MAE is {evaluation.mae:.2f} KSh/L and RMSE is "
        f"{evaluation.rmse:.2f} KSh/L. The intended final design is June 2026 "
        "components to July 2026 prices; however, verified June components are not "
        "available in the repository. July predictions are therefore not published, "
        "preventing target leakage and fabricated accuracy. The Streamlit system "
        "still provides source-linked reconstruction, factor explanation, scenarios, "
        "and budget and journey calculators.",
    )
    add_para(
        doc,
        "Keywords: fuel prices, Nairobi, multiple linear regression, components, "
        "EPRA, Streamlit, fuel budgeting",
        italic=True,
    )

    doc.add_page_break()
    add_heading(doc, "TABLE OF CONTENTS", 1)
    add_para(
        doc,
        "Chapter One: Introduction\nChapter Two: Literature Review\n"
        "Chapter Three: Methodology and System Design\n"
        "Chapter Four: Implementation, Results and Verification\n"
        "Chapter Five: Summary, Conclusions and Recommendations\n"
        "References\nAppendices",
    )
    add_heading(doc, "LIST OF ABBREVIATIONS", 1)
    add_table(
        doc,
        ["Abbreviation", "Meaning"],
        [
            ["EPRA", "Energy and Petroleum Regulatory Authority"],
            ["KES / KSh", "Kenya shilling"],
            ["MAE", "Mean absolute error"],
            ["RMSE", "Root mean squared error"],
            ["MLR", "Multiple linear regression"],
            ["PMS", "Super Petrol"],
            ["AGO", "Automotive Diesel"],
            ["IK / DPK", "Kerosene"],
            ["UI", "User interface"],
        ],
        [1.6, 4.9],
    )

    doc.add_page_break()
    add_heading(doc, "CHAPTER ONE: INTRODUCTION", 1)
    add_heading(doc, "1.1 Background of the Study", 2)
    add_para(
        doc,
        "Fuel is a major operating input for ride-hailing, taxi, matatu, courier, "
        "and logistics work in Nairobi. EPRA publishes maximum retail petroleum "
        "prices for stated effective periods. These are regulatory ceilings rather "
        "than guaranteed station-level selling prices. The final price is influenced "
        "by imported-product cost, distribution, approved margins, taxes, levies, "
        "and stabilization decisions.",
    )
    add_heading(doc, "1.2 Problem Statement", 2)
    add_para(
        doc,
        "Official notices provide authoritative prices but do not provide one simple "
        "workflow for studying component effects, producing a leakage-safe academic "
        "prediction, validating historical build-ups, and converting prices into "
        "fuel budgets and journey costs. Earlier project complexity also made the "
        "method difficult to explain. A simpler, evidence-led system is required.",
    )
    add_heading(doc, "1.3 Main Objective", 2)
    add_para(
        doc,
        "To design and implement a component-based machine-learning system using "
        "multiple linear regression to predict July 2026 maximum retail fuel prices "
        "in Nairobi and support fuel-budget planning.",
    )
    add_heading(doc, "1.4 Specific Objectives", 2)
    add_numbered(
        doc,
        [
            "To collect and validate historical Nairobi fuel-price component data from official sources.",
            "To analyse how landed cost, distribution and storage, margins, stabilization, taxes and levies affect retail fuel prices.",
            "To develop a multiple linear regression model using verified pre-target component information.",
            "To predict July 2026 prices for Super Petrol, Diesel and Kerosene.",
            "To compare predicted July prices with official July prices.",
            "To implement a Streamlit application for prediction, price explanation, reconstruction, and fuel-cost planning.",
        ],
    )
    add_heading(doc, "1.5 Research Questions", 2)
    add_numbered(
        doc,
        [
            "Which cost components affect Nairobi retail fuel prices?",
            "What relationship exists between fuel-price cost components and final retail prices?",
            "How accurately can multiple linear regression predict July 2026 fuel prices?",
            "How can the system support budgeting and journey-cost planning for Nairobi transport users?",
        ],
    )
    add_heading(doc, "1.6 Scope and Users", 2)
    add_para(
        doc,
        "The geographical scope is Nairobi and the products are Super Petrol, "
        "Diesel, and Kerosene. MafutaPlan is designed primarily for Nairobi "
        "transport operators and small transport businesses that need fuel-price "
        "information for budgeting and journey-cost planning.",
    )
    add_para(
        doc,
        "The primary persona is Brian, a Nairobi ride-hailing driver. Brian views "
        "the July prediction status, examines price factors, estimates weekly fuel "
        "expenses, calculates journey costs, and plans his transport budget.",
    )
    add_heading(doc, "1.7 Significance", 2)
    add_para(
        doc,
        "The project demonstrates explainable machine learning, official-source "
        "governance, and a practical application relevant to Nairobi transport. "
        "EPRA is the regulator, authoritative data source, and stakeholder; it is "
        "not described as the project client. This is an academic project designed "
        "for target users rather than a commissioned system.",
    )
    add_heading(doc, "1.8 Limitations", 2)
    add_bullets(
        doc,
        [
            "July components cannot be used to predict July without leakage.",
            "The verified component panel is small and discontinuous.",
            "Official regulatory, tax, and stabilization decisions can change abruptly.",
            "The application does not replace EPRA.",
            "Station-level prices may differ from the regulatory maximum.",
            "Small-sample coefficients and errors have limited generalisability.",
        ],
    )

    doc.add_page_break()
    add_heading(doc, "CHAPTER TWO: LITERATURE REVIEW", 1)
    add_heading(doc, "2.1 Regulated Fuel Pricing", 2)
    add_para(
        doc,
        "EPRA's published pump-price formula states that maximum retail prices are "
        "formed from wholesale price, secondary transport, retail investment and "
        "operating margins, and VAT, with underlying import, storage, transport, "
        "loss, tax, levy, and approved-cost factors. This supports an additive "
        "component representation while recognising that official definitions are "
        "more detailed than the five analytical groups.",
    )
    add_heading(doc, "2.2 Aggregate Cost Groups", 2)
    add_table(
        doc,
        ["Group", "Coverage in MafutaPlan"],
        [
            ["Landed cost", "Product cost, freight, premium, insurance, finance, port/import costs and exchange effects already reflected in landed value."],
            ["Distribution and storage", "Jetty handling, storage, pipeline, allowable losses, depot handling and Nairobi delivery."],
            ["Margins", "Wholesale, retail investment and retail operating margins."],
            ["Taxes and levies", "Excise, VAT, road, development, regulatory and other applicable statutory charges."],
            ["Stabilization adjustment", "Signed subsidy, compensation, deficit, surplus, approved adjustment or rounding reconciliation."],
        ],
        [2.0, 4.5],
    )
    add_para(
        doc,
        "An aggregate group and its detailed subcomponents are not used together in "
        "the model. This avoids double-counting.",
    )
    add_heading(doc, "2.3 Multiple Linear Regression", 2)
    add_para(
        doc,
        "Multiple linear regression estimates a continuous target as an intercept "
        "plus weighted input variables. It is appropriate here because it is "
        "transparent, deterministic, and easy to explain. A positive coefficient "
        "means the fitted prediction generally rises as that input rises, holding "
        "the other included inputs constant. Coefficients are associations within "
        "the observed sample rather than proof of causation.",
    )
    add_heading(doc, "2.4 Evaluation and Leakage", 2)
    add_para(
        doc,
        "Random train-test splitting can allow later regulatory regimes to influence "
        "earlier predictions. MafutaPlan orders records chronologically and reserves "
        "the latest complete target cycle. July 2026 is never included in training. "
        "MAE reports the average absolute error and RMSE gives greater weight to "
        "larger errors.",
    )
    add_heading(doc, "2.5 Conceptual Framework", 2)
    add_figure(
        doc,
        DIAGRAMS / "conceptual_framework.png",
        "Figure 2.1: Conceptual framework for component-based prediction and planning",
    )

    doc.add_page_break()
    add_heading(doc, "CHAPTER THREE: METHODOLOGY AND SYSTEM DESIGN", 1)
    add_heading(doc, "3.1 Research and Development Approach", 2)
    add_para(
        doc,
        "The project uses a design-and-implementation approach. Official EPRA "
        "records are collected, structured, validated, transformed into one-cycle-"
        "ahead model rows, fitted with pooled multiple linear regression, evaluated "
        "chronologically, and exposed through a Streamlit interface.",
    )
    add_heading(doc, "3.2 Data Sources and Integrity", 2)
    add_para(
        doc,
        f"The reviewed component history contains {len(components)} fuel-cycle "
        f"records across {components['Effective_From'].nunique()} official Annex "
        "cycles. Each row records effective dates, fuel type, five aggregate "
        "components, published price, source title, official HTTPS link, verification "
        "status, reconstructed price, and reconstruction error. Missing components "
        "are not imputed or fabricated.",
    )
    add_figure(
        doc,
        CHARTS / "figure_3_1_component_history.png",
        "Figure 3.1: Verified aggregate component history",
    )
    add_heading(doc, "3.3 Model-Ready Dataset", 2)
    add_para(
        doc,
        "For every verified input cycle, the five components and fuel type are "
        "paired with the following cycle's retail price. This creates 33 usable "
        "rows. Gaps in official component coverage remain gaps; the rows are not "
        "presented as a continuous monthly panel.",
    )
    add_table(
        doc,
        ["Variable", "Role"],
        [
            ["Landed_Cost", "Independent variable"],
            ["Distribution_Storage", "Independent variable"],
            ["Margins", "Independent variable"],
            ["Stabilization_Adjustment", "Independent variable"],
            ["Taxes_Levies", "Independent variable"],
            ["Fuel", "Categorical input encoded with Super Petrol as the reference"],
            ["Target_Retail_Price", "Dependent variable: following-cycle Nairobi maximum retail price"],
        ],
        [2.3, 4.2],
    )
    add_heading(doc, "3.4 Model Formula", 2)
    add_para(
        doc,
        "Predicted price = intercept + b1 x landed cost + b2 x distribution and "
        "storage + b3 x margins + b4 x stabilization + b5 x taxes and levies + "
        "fuel-type effect.",
    )
    add_heading(doc, "3.5 Chronological Evaluation Design", 2)
    add_para(
        doc,
        f"Training targets run from {evaluation.training_start:%B %Y} to "
        f"{evaluation.training_end:%B %Y} using {evaluation.training_records} rows. "
        f"The chronological test target is {evaluation.test_cycle:%B %Y} with "
        f"{evaluation.test_records} rows. July 2026 is outside training and is "
        "reserved for final evaluation only when verified June component inputs exist.",
    )
    add_heading(doc, "3.6 Separation of Modules", 2)
    add_table(
        doc,
        ["Module", "Inputs", "Method", "Output"],
        [
            ["Prediction", "Pre-target components and fuel type", "Multiple linear regression", "Following-cycle estimate"],
            ["Reconstruction", "Known same-cycle official components", "Deterministic addition", "Reconstructed official price"],
            ["Scenario", "User-adjusted assumptions", "Deterministic addition", "What-if estimate"],
            ["Calculator", "Price, litres, budget, distance, efficiency, contingency", "Arithmetic formulas", "Purchase and journey costs"],
        ],
        [1.2, 2.1, 1.5, 1.7],
    )
    add_heading(doc, "3.7 System Architecture", 2)
    add_figure(
        doc,
        DIAGRAMS / "system_architecture_diagram.png",
        "Figure 3.2: MafutaPlan system architecture and July availability gate",
    )
    add_heading(doc, "3.8 Use Cases", 2)
    add_figure(
        doc,
        DIAGRAMS / "use_case_diagram.png",
        "Figure 3.3: Brian's MafutaPlan use cases",
        width=5.8,
    )

    doc.add_page_break()
    add_heading(doc, "CHAPTER FOUR: IMPLEMENTATION, RESULTS AND VERIFICATION", 1)
    add_heading(doc, "4.1 Implementation", 2)
    add_para(
        doc,
        "The implementation uses Python, pandas, scikit-learn, and Streamlit. "
        "The application has six pages: Home, July 2026 Prediction, Factors "
        "Affecting Fuel Price, Price Reconstruction, Fuel Calculator, and Data and "
        "Methodology. The model code directly creates the design matrix, fits one "
        "LinearRegression estimator, and reports coefficients and chronological "
        "errors without model-selection abstraction.",
    )
    add_heading(doc, "4.2 Reconstruction Result", 2)
    audit = reconstruction_audit(components)
    add_para(
        doc,
        f"All {len(audit)} reviewed rows reconstruct the stored official retail "
        f"price within the KSh 0.02 validation tolerance. The maximum absolute "
        f"calculated error is {audit['Calculated_Error'].abs().max():.2f} KSh/L.",
    )
    add_heading(doc, "4.3 Learned Coefficients", 2)
    coefficient_rows = [
        [COMPONENT_LABELS.get(row.Term, row.Term), f"{float(row.Coefficient):.6f}"]
        for row in evaluation.coefficients.itertuples(index=False)
    ]
    add_table(doc, ["Term", "Coefficient"], coefficient_rows, [3.8, 2.7])
    add_figure(
        doc,
        CHARTS / "figure_4_2_coefficients.png",
        "Figure 4.1: Fitted coefficients from the chronological training sample",
    )
    add_para(
        doc,
        "The distribution/storage and margin coefficients are unstable in magnitude "
        "because those values vary within narrow bands and are correlated with cycle "
        "and fuel structure in a small sample. They should not be interpreted as "
        "causal pass-through estimates.",
    )
    add_heading(doc, "4.4 Chronological Test Results", 2)
    result_rows = [
        [
            row.Fuel,
            f"{row.Target_Retail_Price:.2f}",
            f"{row.Predicted_Retail_Price:.2f}",
            f"{row.Absolute_Error:.2f}",
            f"{row.Percentage_Error:.2f}%",
        ]
        for row in evaluation.results.itertuples(index=False)
    ]
    add_table(
        doc,
        ["Fuel", "Actual", "Predicted", "Absolute error", "Percentage error"],
        result_rows,
        [1.4, 1.1, 1.2, 1.4, 1.4],
    )
    add_para(
        doc,
        f"Overall chronological test MAE = {evaluation.mae:.2f} KSh/L; "
        f"RMSE = {evaluation.rmse:.2f} KSh/L. The abrupt April 2026 price changes "
        "produce large errors for Super Petrol and Diesel, demonstrating the effect "
        "of regulatory decisions that are not fully learnable from the small component panel.",
    )
    add_figure(
        doc,
        CHARTS / "figure_4_1_actual_vs_predicted.png",
        "Figure 4.2: Actual versus predicted April 2026 holdout prices",
    )
    add_heading(doc, "4.5 July 2026 Final Evaluation Status", 2)
    official_row = official.iloc[0]
    add_table(
        doc,
        ["Fuel", "Official July price", "Prediction", "Error"],
        [
            ["Super Petrol", f"{float(official_row.Super_Petrol):.2f}", "Unavailable", "Unavailable"],
            ["Diesel", f"{float(official_row.Diesel):.2f}", "Unavailable", "Unavailable"],
            ["Kerosene", f"{float(official_row.Kerosene):.2f}", "Unavailable", "Unavailable"],
        ],
        [1.7, 1.6, 1.6, 1.6],
    )
    add_para(
        doc,
        "Verified June 2026 component inputs are absent. July predictions, "
        "absolute errors, percentage errors, MAE, and RMSE are therefore not "
        "claimed. July official values are retained only as final evaluation values. "
        "Using July components would be target leakage; substituting March components "
        "would change the one-cycle-ahead horizon without evidence.",
    )
    doc.add_page_break()
    add_heading(doc, "CHAPTER FIVE: SUMMARY, CONCLUSIONS AND RECOMMENDATIONS", 1)
    add_heading(doc, "5.1 Summary", 2)
    add_para(
        doc,
        "MafutaPlan was simplified into a component-based academic system with one "
        "pooled multiple linear regression model. The model uses five non-overlapping "
        "aggregate cost groups and encoded fuel type. Historical reconstruction, "
        "scenario analysis, and calculators remain deterministic and separate.",
    )
    add_heading(doc, "5.2 Conclusions", 2)
    add_para(
        doc,
        "The architecture is valid and explainable, but the available evidence is "
        "insufficient for a defensible July 2026 prediction. The high April holdout "
        "error and missing June inputs show why verified data and clear limitations "
        "are more important than manufacturing a successful forecast. The system's "
        "practical value also comes from price explanation, reconstruction, and "
        "transport budgeting.",
    )
    add_heading(doc, "5.3 Recommendations", 2)
    add_numbered(
        doc,
        [
            "Obtain and verify the official June 2026 component annex before generating July predictions.",
            "Expand the component panel with continuous official cycles without interpolation.",
            "Retain chronological evaluation and keep July outside training.",
            "Treat stabilization and tax decisions as policy-sensitive limitations.",
            "Conduct structured usability testing with Nairobi transport operators.",
            "Refresh official sources and revision records without overwriting audit history.",
        ],
    )

    doc.add_page_break()
    add_heading(doc, "REFERENCES", 1)
    add_bullets(
        doc,
        [
            "Energy and Petroleum Regulatory Authority. Pump Price Formulae. https://www.epra.go.ke/pump-price-formulae",
            "Energy and Petroleum Regulatory Authority. Pump Prices. https://www.epra.go.ke/pump-prices",
            "Energy and Petroleum Regulatory Authority. Official component annex releases listed in data/sources.csv.",
            "Pedregosa, F., et al. (2011). Scikit-learn: Machine Learning in Python. Journal of Machine Learning Research, 12, 2825-2830.",
            "Streamlit. Documentation. https://docs.streamlit.io/",
        ],
    )

    add_heading(doc, "APPENDICES", 1)
    add_heading(doc, "Appendix A: Data Dictionary", 2)
    dictionary = pd.read_csv(APPENDICES / "Data_Extraction_Sheet.csv")
    add_table(
        doc,
        dictionary.columns.tolist(),
        dictionary.astype(str).values.tolist(),
        [1.2, 0.75, 0.75, 2.5, 1.3],
    )
    add_heading(doc, "Appendix B: Source Register", 2)
    source_rows = sources[["Source_ID", "Publisher", "Title"]].astype(str).values.tolist()
    add_table(doc, ["Source ID", "Publisher", "Title"], source_rows, [1.6, 2.0, 2.9])
    add_para(
        doc,
        "Complete HTTPS links, access dates, and provenance notes are retained in "
        "data/sources.csv.",
        italic=True,
    )
    add_heading(doc, "Appendix C: User Guide", 2)
    add_numbered(
        doc,
        [
            "Install runtime dependencies with python -m pip install -r requirements.txt.",
            "Start the application with streamlit run app.py.",
            "Use Home for scope, users, persona, and official July evaluation values.",
            "Use July 2026 Prediction to inspect the input-availability decision.",
            "Use Factors Affecting Fuel Price for component shares and deterministic scenarios.",
            "Use Price Reconstruction to reproduce a verified historical price and open its source.",
            "Use Fuel Calculator for purchase, budget, and journey calculations.",
            "Use Data and Methodology for coefficients, MAE, RMSE, results, sources, and limitations.",
        ],
    )
    add_heading(doc, "Appendix D: Reproduction Commands", 2)
    add_bullets(
        doc,
        [
            "python scripts/build_model_dataset.py",
            "python scripts/build_notebook.py",
            "python scripts/build_report.py",
            "python -m compileall app.py src scripts",
            "python -m pip check",
            "streamlit run app.py",
        ],
    )
    add_heading(doc, "Appendix E: Data Gaps", 2)
    add_para(
        doc,
        "The verified component cycles are August, October, November, and December "
        "2024; February, June, July, and August 2025; and January, February, and "
        "March 2026. Other months, including June 2026, are not included as verified "
        "component records. No missing record has been interpolated.",
    )
    return doc


def main() -> None:
    DOCS.mkdir(parents=True, exist_ok=True)
    components = load_component_history()
    prediction_data = load_prediction_dataset()
    sources = load_sources()
    official = load_official_prices()
    evaluation = evaluate_latest_cycle(prediction_data)
    build_figures(components, evaluation)
    write_appendices(prediction_data, evaluation)
    report = build_report(components, prediction_data, sources, official, evaluation)
    report.save(REPORT_PATH)
    print(f"Wrote {REPORT_PATH.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
